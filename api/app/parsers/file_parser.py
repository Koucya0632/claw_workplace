from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    # ParsedDocument 是 parser 輸出的統一格式，供索引服務直接落庫。
    text: str
    preview: str
    mime_type: str
    extension: str


class FileParser:
    # FileParser 依副檔名選擇不同解析器，Phase 1 專注於文本抽取。
    MIME_MAP = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
        "md": "text/markdown",
        "csv": "text/csv",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }

    def parse(self, path: Path, payload: bytes) -> ParsedDocument:
        # 副檔名一律標準化為小寫，避免大小寫造成 parser 選錯。
        extension = path.suffix.lower().lstrip(".")

        if extension == "pdf":
            text = self._parse_pdf(payload)
        elif extension == "docx":
            text = self._parse_docx(payload)
        elif extension in {"txt", "md"}:
            text = self._parse_text(payload)
        elif extension == "csv":
            text = self._parse_csv(payload)
        elif extension == "xlsx":
            text = self._parse_xlsx(payload)
        else:
            raise ValueError(f"不支援的副檔名：{extension}")

        # 把連續空白壓掉，讓預覽與摘要輸入更乾淨。
        normalized_text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        preview = normalized_text[:320]

        return ParsedDocument(
            text=normalized_text,
            preview=preview,
            mime_type=self.MIME_MAP.get(extension, "application/octet-stream"),
            extension=extension,
        )

    def _parse_text(self, payload: bytes) -> str:
        # 純文字與 Markdown 直接用 UTF-8 讀取，遇到壞字元時用 replace 保留最大資訊量。
        return payload.decode("utf-8", errors="replace")

    def _parse_csv(self, payload: bytes) -> str:
        # CSV 轉成以 tab 分隔的文字，讓表格內容也能被全文搜索命中。
        decoded = payload.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(decoded))
        lines = ["\t".join(cell.strip() for cell in row) for row in reader]
        return "\n".join(lines)

    def _parse_xlsx(self, payload: bytes) -> str:
        # Excel 每張工作表都轉成純文字段落，保留工作表名稱幫助理解來源。
        workbook = load_workbook(io.BytesIO(payload), read_only=True, data_only=True)
        lines: list[str] = []
        for worksheet in workbook.worksheets:
            lines.append(f"[Sheet] {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    lines.append("\t".join(values))
        return "\n".join(lines)

    def _parse_docx(self, payload: bytes) -> str:
        # Word 檔逐段抽出文字內容，忽略空段落。
        document = Document(io.BytesIO(payload))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    def _parse_pdf(self, payload: bytes) -> str:
        # PDF 使用 pypdf 做簡單文字抽取，足夠支援 Phase 1 索引與摘要。
        reader = PdfReader(io.BytesIO(payload))
        pages: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                pages.append(extracted.strip())
        return "\n".join(pages)

